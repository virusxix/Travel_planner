"""
Rebuild C.8 Gantt Chart in HiddenStay_v4_Final.docx to match sample format:
- Phase legend with color keys
- Task rows with assignees in parentheses
- Horizontal colored bars across W1–W10 + date ranges
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DOC_PATH = Path(r"d:\sample\HiddenStay_v4_Final.docx")
OUT_PATH = Path(r"d:\sample\HiddenStay_v4_Final.docx")

# Phase legend colors (sample-style palette)
PHASES = {
    "Team Formation & Initiation": "F687B3",
    "Project Planning & Requirements": "F6E05E",
    "Design & Architecture": "68D391",
    "Backend Development": "A0AEC0",
    "Frontend Development": "63B3ED",
    "AI & Maps Integration": "9F7AEA",
    "Testing & QA": "B2F5EA",
    "Deploy & Submission": "FBD38D",
}

WEEK_DATES = [
    "Jun 3–6",
    "Jun 9–13",
    "Jun 16–20",
    "Jun 23–27",
    "Jun 30–Jul 4",
    "Jul 7–11",
    "Jul 14–18",
    "Jul 21–25",
    "Jul 28–Aug 1",
    "Aug 4–8",
]

# (phase, task label with assignee, active week indices 1-10)
TASKS = [
    ("Team Formation & Initiation", "Set charter, repo & team norms (All members)", [1]),
    ("Team Formation & Initiation", "Literature review & OTA commission analysis (Business Analyst)", [1, 2]),
    ("Team Formation & Initiation", "Competitive landscape & commission model validation (Business Analyst)", [2]),
    ("Project Planning & Requirements", "Define functional requirements & user stories (PM + Business Analyst)", [2, 3]),
    ("Project Planning & Requirements", "System architecture & ERD design (Lead Developer)", [3]),
    ("Project Planning & Requirements", "Complete project proposal draft (All members)", [3, 4]),
    ("Design & Architecture", "Wireframes & traveller/host user journeys (UI/UX Designer)", [3, 4]),
    ("Design & Architecture", "UI/UX design — web screens & brand identity (UI/UX Designer)", [4, 5]),
    ("Backend Development", "Prisma schema & Supabase database setup (Lead Developer)", [4, 5]),
    ("Backend Development", "JWT auth, property & room listing API (Lead Developer)", [5, 6]),
    ("Frontend Development", "Landing page, search & auth UI (UI/UX Designer + Lead Developer)", [5, 6]),
    ("Frontend Development", "Property detail & booking UI shell (Lead Developer)", [6]),
    ("Backend Development", "Stripe sandbox & booking confirmation flow (Lead Developer)", [6, 7]),
    ("Backend Development", "Hotel owner portal & admin approval workflow (Lead Developer)", [7]),
    ("Team Formation & Initiation", "Pilot hotel outreach — target 5 listings (Marketing & Partnerships Lead)", [7, 8]),
    ("AI & Maps Integration", "AI itinerary JSON engine — Groq/Gemini/OpenAI (Lead Developer)", [7, 8]),
    ("AI & Maps Integration", "Google Places geocoding & map route optimisation (Lead Developer)", [7, 8]),
    ("AI & Maps Integration", "Prompt engineering & itinerary UI pages (PM + Business Analyst)", [7, 8, 9]),
    ("AI & Maps Integration", "AI travel chat & dashboard CRUD integration (Lead Developer)", [8, 9]),
    ("Testing & QA", "Unit, integration & UAT with pilot partners (All members)", [9]),
    ("Testing & QA", "Bug fixes & performance validation (Lead Developer)", [9]),
    ("Deploy & Submission", "Staging deployment & security checklist (Lead Developer + PM)", [10]),
    ("Deploy & Submission", "Final documentation & capstone submission (All members)", [10]),
]

HEADER_NAVY = "1A365D"
LABEL_GRAY = "EDF2F7"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.find(qn("w:shd"))
    if old is not None:
        tc_pr.remove(old)
    tc_pr.append(shading)


def set_cell_text(cell, text: str, *, bold=False, color=None, size=8, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color


def delete_table(table):
    table._element.getparent().remove(table._element)


def find_gantt_tables(doc: Document):
    gantt = legend = None
    for table in doc.tables:
        if not table.rows:
            continue
        header = [c.text.strip() for c in table.rows[0].cells]
        joined = " ".join(header)
        if header[:3] == ["Phase / Track", "W1", "W2"]:
            gantt = table
        elif "Lead Dev Track" in joined or "All-team Phase" in joined:
            legend = table
    return gantt, legend


def build_legend_table(doc: Document):
    """2-row legend: phase name + color swatch, sample-style."""
    rows_data = list(PHASES.items())
    table = doc.add_table(rows=2, cols=len(rows_data))

    for i, (name, color) in enumerate(rows_data):
        swatch = table.rows[0].cells[i]
        label = table.rows[1].cells[i]
        set_cell_shading(swatch, color)
        set_cell_text(swatch, "", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(label, name, bold=True, size=7, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(label, WHITE)

    return table


def build_gantt_table(doc: Document):
    cols = 11  # task label + W1..W10
    table = doc.add_table(rows=2 + len(TASKS), cols=cols)

    # Header row 0: Task | W1..W10
    set_cell_text(table.rows[0].cells[0], "Task (Assignee)", bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(table.rows[0].cells[0], HEADER_NAVY)
    for w in range(1, 11):
        set_cell_text(table.rows[0].cells[w], f"W{w}", bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[w], HEADER_NAVY)

    # Header row 1: dates
    set_cell_text(table.rows[1].cells[0], "", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(table.rows[1].cells[0], LABEL_GRAY)
    for w in range(1, 11):
        set_cell_text(table.rows[1].cells[w], WEEK_DATES[w - 1], bold=False, size=7, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[1].cells[w], LABEL_GRAY)

    # Task rows
    for ri, (phase, task, weeks) in enumerate(TASKS, start=2):
        color = PHASES[phase]
        row = table.rows[ri]
        set_cell_text(row.cells[0], task, size=8)
        set_cell_shading(row.cells[0], LABEL_GRAY)
        week_set = set(weeks)
        for w in range(1, 11):
            cell = row.cells[w]
            if w in week_set:
                set_cell_shading(cell, color)
                set_cell_text(cell, "", align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                set_cell_shading(cell, WHITE)
                set_cell_text(cell, "", align=WD_ALIGN_PARAGRAPH.CENTER)

    # Column widths
    for row in table.rows:
        row.cells[0].width = Inches(2.8)
        for w in range(1, 11):
            row.cells[w].width = Inches(0.45)

    return table


def insert_after_paragraph(paragraph, element):
    paragraph._element.addnext(element)


def main():
    doc = Document(DOC_PATH)

    # Update intro paragraph
    for p in doc.paragraphs:
        if p.text.strip().startswith("The chart below illustrates the 10-week"):
            p.clear()
            run = p.add_run(
                "Figure C.8 — HiddenStay AI 10-Week Capstone Gantt Chart. "
                "Each row represents an individual task with the responsible team member(s) noted in parentheses. "
                "Coloured bars indicate active work periods across Weeks 1–10 (Jun 3 – Aug 8, 2026). "
                "Parallel tracks — particularly AI integration and backend development during Weeks 7–9 — "
                "are shown as overlapping bars on separate rows, following standard capstone project scheduling practice."
            )
            run.font.size = Pt(10)
            break

    gantt_old, legend_old = find_gantt_tables(doc)

    # Anchor: paragraph before old gantt table
    anchor = None
    for p in doc.paragraphs:
        if "Figure C.8" in p.text or p.text.strip().startswith("The chart below illustrates"):
            anchor = p
            break
    if anchor is None:
        for p in doc.paragraphs:
            if p.text.strip() == "C.8   Gantt Chart":
                anchor = p
                break

    # Remove old tables
    if legend_old is not None:
        delete_table(legend_old)
    if gantt_old is not None:
        delete_table(gantt_old)

    # Build new tables and insert after anchor paragraph
    legend_tbl = build_legend_table(doc)
    gantt_tbl = build_gantt_table(doc)

    if anchor is not None:
        # Insert gantt first, then legend after anchor — final order: anchor → legend → gantt
        insert_after_paragraph(anchor, gantt_tbl._element)
        insert_after_paragraph(anchor, legend_tbl._element)
    else:
        doc.add_paragraph("C.8 Gantt Chart (rebuilt)")

    try:
        doc.save(OUT_PATH)
        saved = OUT_PATH
    except PermissionError:
        saved = OUT_PATH.with_name("HiddenStay_v4_Final_GanttFixed.docx")
        doc.save(saved)
        print("NOTE: Original file locked (close Word). Saved alternate copy.")

    print(f"Sample-style Gantt rebuilt: {saved}")
    print(f"  Legend phases: {len(PHASES)}")
    print(f"  Task rows: {len(TASKS)}")


if __name__ == "__main__":
    main()
