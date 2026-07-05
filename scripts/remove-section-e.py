"""Remove Section E from HiddenStay_v4_Final.docx and fix cross-references."""
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

DOC_PATH = Path(r"d:\sample\HiddenStay_v4_Final.docx")
OUT_PATH = DOC_PATH

B3_TECH_NOTE = (
    "Technical approach (AI module): The AI Trip Planner will use a structured JSON prompt pipeline, "
    "Google Places enrichment for geocoded venue coordinates, and a multi-provider LLM fallback chain "
    "(Groq, Gemini, OpenAI) to ensure reliable itinerary generation during the capstone demonstration period."
)


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def delete_table(table):
    table._element.getparent().remove(table._element)


def is_section_e_heading(text: str) -> bool:
    t = text.strip()
    return (
        t.startswith("E.   AI Trip Planner")
        or t.startswith("E.1")
        or t.startswith("E.2")
        or t.startswith("E.3")
        or t.startswith("E.4")
        or t.startswith("E.5")
    )


def is_section_e_toc(text: str) -> bool:
    t = text.strip()
    return t.startswith("E.   AI Trip Planner") or (
        t.startswith("E.") and "\t" in t and any(x in t for x in ("15", "16", "17"))
    )


def is_section_e_body_paragraph(text: str, in_section_e: bool) -> bool:
    if not in_section_e:
        return False
    t = text.strip()
    if t.startswith("F.   References"):
        return False
    if is_section_e_heading(t):
        return True
    if t.startswith("F."):
        return False
    return True


def table_is_section_e(table) -> bool:
    full = " ".join(c.text for row in table.rows for c in row.cells)
    if "Layer" in full and "Technology" in full and "Next.js" in full:
        return True
    if "Stage 1" in full and "Soft Delete" in full:
        return True
    if "Primary:" in full and "Groq" in full and "Fallback" in full:
        return True
    if "Success Metric" in full and "Platform deployment (staging" in full:
        return True
    return False


def replace_in_cell(cell, old: str, new: str):
    for paragraph in cell.paragraphs:
        if old in paragraph.text:
            for run in paragraph.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)


def replace_cell_text(cell, replacements: list[tuple[str, str]]):
    full = cell.text
    for old, new in replacements:
        full = full.replace(old, new)
    if full != cell.text:
        cell.text = full


def fix_cross_references(doc: Document):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_cell_text(cell, [
                    (
                        "Deploy 5 verified pilot listings for the 10-week Capstone evaluation phase "
                        "(see Section E.5), laying the infrastructure",
                        "Deploy 5 verified pilot listings for the 10-week Capstone evaluation phase, "
                        "laying the infrastructure",
                    ),
                    ("(see Section E.5)", ""),
                    ("see Section E.5", "see Section B.5"),
                    (
                        "Soft Delete workflow for user data (see Section E.3).",
                        "Soft-delete workflow for user itinerary data.",
                    ),
                    ("soft-delete workflow verified (Section E.3)", "soft-delete workflow verified"),
                    ("see Section E.3", ""),
                ])

    for p in doc.paragraphs:
        if "see Section E" in p.text:
            for run in p.runs:
                run.text = run.text.replace(" (see Section E.3 for soft-delete workflow)", "")
                run.text = run.text.replace("see Section E.3 for soft-delete workflow", "soft-delete workflow")


def add_b3_technical_note(doc: Document):
    if any("Technical approach (AI module)" in p.text for p in doc.paragraphs):
        return

    target = None
    for p in doc.paragraphs:
        if p.text.strip() == "Deployment to a publicly accessible staging environment (Vercel + Supabase)":
            target = p
            break
    if target is None:
        return

    # Insert after deployment bullet
    new_p = target._element
    parent = new_p.getparent()
    idx = list(parent).index(new_p)

    from docx.oxml import OxmlElement

    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = B3_TECH_NOTE
    r.append(t)
    p.append(r)
    parent.insert(idx + 1, p)


def remove_toc_section_e(doc: Document):
    for p in list(doc.paragraphs):
        t = p.text.strip()
        if is_section_e_toc(t):
            delete_paragraph(p)


def remove_body_section_e(doc: Document):
    """Remove all Section E content between D (stakeholders) and F (references)."""
    d_body_idx = None
    f_body_idx = None

    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "D.   Stakeholder Analysis":
            d_body_idx = i
        if t == "F.   References" and d_body_idx is not None:
            f_body_idx = i
            break

    if d_body_idx is None or f_body_idx is None:
        # Fallback: marker-based deletion
        to_delete = []
        in_e = False
        for p in doc.paragraphs:
            t = p.text.strip()
            if t.startswith("E.   AI Trip Planner") or t.startswith("This section documents the technical architecture"):
                in_e = True
            if in_e and t.startswith("F.   References"):
                in_e = False
                continue
            if in_e or t.startswith("E.1") or t.startswith("E.2") or t.startswith("E.3") or t.startswith("E.4") or t.startswith("E.5"):
                to_delete.append(p)
        for p in to_delete:
            delete_paragraph(p)
        return

    # Delete only Section E paragraphs between D and F (preserve D intro + stakeholder table)
    to_delete = []
    for i in range(d_body_idx + 1, f_body_idx):
        t = doc.paragraphs[i].text.strip()
        if not t:
            continue
        if t.startswith("A stakeholder analysis"):
            continue
        to_delete.append(doc.paragraphs[i])

    for p in to_delete:
        delete_paragraph(p)


def remove_section_e_tables(doc: Document):
    for table in list(doc.tables):
        if table_is_section_e(table):
            delete_table(table)


def main():
    doc = Document(DOC_PATH)

    remove_toc_section_e(doc)
    remove_body_section_e(doc)
    remove_section_e_tables(doc)
    fix_cross_references(doc)
    add_b3_technical_note(doc)

    try:
        doc.save(OUT_PATH)
        saved = OUT_PATH
    except PermissionError:
        saved = OUT_PATH.with_name("HiddenStay_v4_Final_NoSectionE.docx")
        doc.save(saved)

    print(f"Saved: {saved}")
    print("Removed: Section E (E.1–E.5), TOC entries, and 4 technical tables")
    print("Fixed: cross-references in Objectives and Risk tables")
    print("Added: B.3 technical approach note")


if __name__ == "__main__":
    main()
