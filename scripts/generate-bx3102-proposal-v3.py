"""Generate BX3102 A1 Proposal v3.0 Word document from structured content."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
OUT_DOCS = ROOT / "docs" / "BX3102_A1_Proposal_v3.docx"
OUT_DOWNLOADS = Path(r"d:\c\Downloads\BX3102_A1_Proposal_v3.docx")


def set_cell_shading(cell, fill: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    return p


def add_table(doc, headers, rows, header_fill="1A365D"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table


def build_document():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # PAGE 1 — Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("HIDDENSTAY AI")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(
        "A Fair-Commission Accommodation Marketplace\n"
        "with Integrated AI Trip Planner"
    )
    sr.font.size = Pt(14)
    sr.italic = True

    doc.add_paragraph()
    add_table(
        doc,
        ["Field", "Detail"],
        [
            ("Course", "BX3102 — Business and Information Technology Capstone"),
            ("Document", "Assignment 1 — Project Proposal"),
            ("Version", "3.0 (Complete Academic & Technical Overhaul)"),
            ("Submission Date", "June 2026"),
            ("Classification", "Confidential — Academic Submission"),
            ("Repository", "github.com/oakkar-cm/Travel"),
        ],
    )

    doc.add_paragraph()
    add_heading(doc, "Prepared by (5-Member Interdisciplinary Squad)", 2)
    add_table(
        doc,
        ["No.", "Member Name", "Track", "Role"],
        [
            ("1", "________________________", "Business", "Project Manager & Systems Documentation Editor"),
            ("2", "________________________", "Business", "User Experience & Marketplace Strategy Planner"),
            ("3", "________________________", "IT", "Lead Full-Stack Frontend UI Developer"),
            ("4", "________________________", "IT", "Backend Middleware & Transactional Engineer"),
            ("5", "________________________", "IT", "AI Orchestration & DevOps Database Administrator"),
        ],
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "Project Title: HiddenStay AI — Lower-Commission OTA Marketplace and Map-Ready AI Trip "
        "Planner for Independent Hospitality Operators in Southeast Asia."
    )
    doc.add_paragraph(
        "Commission Model (Engine-Enforced): 5% base host fee (Capstone MVP); "
        "10–15% promotional tier (Future Vision)."
    )

    doc.add_page_break()

    # PAGE 2 — TOC & Problem
    add_heading(doc, "PAGE 2 — TABLE OF CONTENTS & PROBLEM STATEMENT", 1)
    add_heading(doc, "Table of Contents", 2)
    for item in [
        "1. Problem Statement — OTA Dependency vs. Independent Margin Leakage",
        "2. SMART Objectives & Viability Validation",
        "3. Project Organisation & Stakeholder Matrix",
        "4. Technical Scope & System Architecture Blueprint",
        "5. Technical Issues Resolution Matrix",
        "6. Management Plan, Conduct & Interdisciplinary Roles",
        "7. Action Plan & Gantt Milestones (10-Week Capstone)",
        "8. Risk Management Matrix",
        "9. Academic Bibliography",
    ]:
        doc.add_paragraph(item, style="List Number")

    add_heading(doc, "A. Problem Statement — OTA Dependency vs. Independent Margin Leakage", 2)
    doc.add_paragraph(
        "Independent and small-scale accommodation operators across Southeast Asia (SEA) are "
        "structurally dependent on dominant Online Travel Agencies (OTAs), principally Booking.com "
        "and Expedia, to sustain online visibility and booking volume. This dependency is not "
        "commercially neutral. Empirical analysis by Juniarta et al. (2026) demonstrates that OTA "
        "distribution channels produce contradictory outcomes: OTA channels exhibit a strong positive "
        "effect on room sales volume (β₁ = 0.897, p < 0.01), yet simultaneously impose a "
        "statistically significant negative effect on hotel profit margins (β₂ = −0.652, p < 0.01). "
        "Economy-class properties bear the most severe consequence, recording an average −SGD 3.47 "
        "EBITDA per available room attributable to disproportionately high commission costs relative "
        "to revenue (Juniarta et al., 2026)."
    )
    doc.add_paragraph(
        "Corporate OTA commission fees typically range from 15% to 30% per completed booking. For "
        "independent operators without chain-level negotiating leverage, this fee structure compresses "
        "operating margins to unsustainable levels. Three reinforcing market conditions exacerbate the problem:"
    )
    for item in [
        "Channel concentration asymmetry: Independent hotels in SEA route approximately 61% of online bookings through OTAs, compared with 35% for branded chain hotels that can negotiate preferential rates (Hospitality Net, 2024).",
        "Commission rate escalation: Expedia and Booking.com have increased fee structures for independent properties, further eroding per-booking retention (Hospitality Technology, 2024; PhocusWire, 2024).",
        "Tier-dependent profitability: Economy and budget properties — the majority of SEA independent inventory — experience net-negative profitability impacts from OTA reliance, whereas luxury-tier operators absorb commission costs through higher average daily rates (Juniarta et al., 2026).",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph(
        "Secondary market failure — fragmented traveller trip planning: Parallel to supply-side margin "
        "leakage, demand-side travellers confront fragmented digital workflows. Property discovery, "
        "itinerary construction, map navigation, and booking confirmation typically occur across "
        "disconnected applications. Generic AI travel assistants frequently output unstructured narrative "
        "itineraries lacking geocoded stop coordinates, verifiable venue identities, or integration with "
        "bookable independent accommodation inventory."
    )
    doc.add_paragraph(
        "Project response: HiddenStay AI addresses both failures within a unified web platform: a 5% "
        "base-commission marketplace for independent stays, coupled with an AI Trip Planner that "
        "generates structured, Google Places–enriched, map-ready itineraries."
    )

    doc.add_page_break()

    # PAGE 3 — SMART Objectives
    add_heading(doc, "PAGE 3 — SMART OBJECTIVES & VIABILITY VALIDATION", 1)
    doc.add_paragraph(
        "All objectives conform to SMART criteria. Objectives are partitioned into Capstone MVP "
        "(10 weeks) and Post-Capstone Commercial Beta."
    )
    add_heading(doc, "B.1 Capstone MVP Objectives (Weeks 1–10)", 2)
    add_table(
        doc,
        ["ID", "Objective", "Success Indicator / KPI", "Measurement Method"],
        [
            ("O1", "Deploy secure web MVP with three role-based portals by Week 10.", "100% core journeys pass UAT.", "UAT script log; assessor checklist."),
            ("O2", "Implement 5% base commission in booking logic.", "5.00% on all sandbox bookings.", "SQL audit; unit tests."),
            ("O3", "Onboard ≥5 verified pilot hotel listings.", "≥5 live listings with full metadata.", "DB property count; quality review."),
            ("O4", "End-to-end Stripe Checkout (Sandbox) booking.", "≥95% success across 5 test bookings.", "Stripe dashboard; error logs."),
            ("O5", "AI itinerary generation with valid JSON plans.", "≥90% of 10 prompts return parseable JSON with ≥3 activities/day.", "Generation log; schema validator."),
            ("O6", "≥80% coordinate accuracy on AI stops.", "≥80% activities have valid Places lat/lng.", "20-itinerary spot-check vs Google Maps."),
            ("O7", "Optimised map routes with branded polylines.", "Route renders; path shorter than chronological sort.", "Haversine comparison; map demo."),
            ("O8", "Traveller dashboard CRUD with soft-delete.", "100% of 5 delete/update scenarios correct DB state.", "API tests; deletedAt audit."),
            ("O9", "Context-aware AI chat with delete coordinator.", "100% of 5 delete intents resolve correctly.", "Chat API validation; unit tests."),
            ("O10", "Complete documentation and staging deploy.", "100% inventory; staging URL live.", "/docs checklist; smoke test."),
        ],
    )
    add_heading(doc, "B.2 Post-Capstone Commercial Beta (Future Vision)", 2)
    add_table(
        doc,
        ["ID", "Objective", "Success Indicator / KPI", "Measurement Method"],
        [
            ("FV1", "Onboard 50 hotel partners in SG and Malaysia.", "50 admin-approved listings.", "Partner onboarding log."),
            ("FV2", "≥500 completed bookings/month by Year 1.", "Monthly volume ≥500.", "Analytics dashboard."),
            ("FV3", "Break-even within 24 months of launch.", "Revenue ≥ costs for 3 consecutive months.", "P&L statements."),
        ],
    )
    add_heading(doc, "C. Why It Is Practical and Necessary", 2)
    doc.add_paragraph(
        "Practical viability: The capstone MVP employs Next.js 15, Express, Prisma, PostgreSQL/Supabase, "
        "Stripe sandbox, Google Maps/Places, and multi-provider LLM endpoints — all demonstrable within "
        "10 weeks by a five-member squad executing parallel workstreams."
    )
    doc.add_paragraph(
        "Economic necessity: At 5% commission on SGD 100/night, a host retains SGD 95.00 versus "
        "SGD 70.00–80.00 under 20–30% OTA fees — a differential that determines viability for "
        "economy operators (Juniarta et al., 2026; Rahadi et al., 2021)."
    )

    doc.add_page_break()

    # PAGE 4 — Organisation
    add_heading(doc, "PAGE 4 — PROJECT ORGANISATION & STAKEHOLDER MATRIX", 1)
    add_heading(doc, "D. Interdisciplinary Product Squad", 2)
    add_table(
        doc,
        ["No.", "Role", "Key Responsibilities", "Skills / Coursework", "Justification"],
        [
            ("1", "PM & Documentation Editor (Business)", "Sprint planning, risk register, report editing", "PM; business communication", "Rubric-aligned deliverables"),
            ("2", "UX & Marketplace Strategy (Business)", "Wireframes, commission modelling, UAT", "HCI; marketing research", "Commercial-UX bridge"),
            ("3", "Lead Frontend Developer (IT)", "Next.js 15, Maps SDK, planner UI, dashboard CRUD", "React/TypeScript", "Highest UI complexity"),
            ("4", "Backend & Transactions (IT)", "Express, Prisma, JWT, Stripe, soft-delete", "API design; security", "Data integrity owner"),
            ("5", "AI & DevOps DBA (IT)", "LLM fallback chain, Places enrichment, Supabase", "ML; cloud computing", "Parallel AI sprint Weeks 7–8"),
        ],
    )
    doc.add_paragraph(
        "Parallel sprint design (Weeks 7–8): Backend finalises booking/dashboard APIs while "
        "AI/DevOps implements itinerary generation; Frontend integrates both into planner UI."
    )
    add_heading(doc, "E. Stakeholder Analysis Matrix", 2)
    add_table(
        doc,
        ["Stakeholder", "Primary Interests", "Expectations", "Capstone Engagement"],
        [
            ("Travellers", "Affordable stays; integrated planning", "Search/book; AI itineraries; dashboard CRUD", "UAT; ≥10 survey respondents"),
            ("Hotel Owners", "Lower commission; visibility", "5% fee; self-service listing", "Owner portal demo"),
            ("Administrators", "Quality; compliance", "Approval workflow; audit logs", "Admin portal demo"),
            ("Project Team", "HD grade; competence", "Clear ownership; deployable artefact", "Stand-ups; weekly demos"),
            ("Assessors", "Rigor; realistic scope", "SMART KPIs; working MVP", "A1 proposal; live demo"),
            ("Investors", "Scalability; ROI", "Pilot metrics; revenue model", "Post-capstone only"),
        ],
    )

    doc.add_page_break()

    # PAGE 5 — Architecture
    add_heading(doc, "PAGE 5 — TECHNICAL SCOPE & SYSTEM ARCHITECTURE", 1)
    add_heading(doc, "F.1 Scope Partition", 2)
    doc.add_paragraph(
        "Part A — Capstone MVP (10 Weeks): Auth, property/room CRUD, search, Stripe sandbox booking, "
        "admin approval, AI planner, map routes, AI chat, soft-delete itineraries, staging deployment."
    )
    doc.add_paragraph(
        "Part B — Future Vision: 50+ hotel beta in SG/MY, mobile apps, 10–15% tiers, six-market expansion, "
        "500+ monthly bookings, 24-month break-even."
    )
    add_heading(doc, "F.2 Four-Tier Production-Grade Web Architecture", 2)
    arch = doc.add_paragraph()
    arch.add_run(
        "[Tier 1: Frontend UI Layer]\n"
        "  Next.js 15 (App Router) · React · TypeScript · Tailwind CSS\n"
        "  Zustand (client UI state) · React Query (server state/cache)\n"
        "        ↕ HTTPS / REST JSON\n"
        "[Tier 2: Backend Middleware Engine Layer]\n"
        "  Node.js · Express API Router · JWT session guards · Zod validation\n"
        "  Prisma ORM query dispatchers · Rate limiting on AI/auth routes\n"
        "        ↕ Prisma Client / SQL\n"
        "[Tier 3: Relational Data Storage Engine]\n"
        "  PostgreSQL on Supabase (users, properties, rooms, bookings,\n"
        "  itineraries, itinerary_days, itinerary_activities — soft-delete columns)\n"
        "        ↕ HTTPS API calls\n"
        "[Tier 4: External Core API Integrations]\n"
        "  AI: Groq (LLaMA-3) → Google Gemini Pro → OpenAI (fallback chain)\n"
        "  Maps: Google Maps JavaScript SDK + Google Places API\n"
        "  Payments: Stripe Checkout API (Sandbox)\n"
        "  Media: Cloudinary (property images)"
    ).font.name = "Consolas"

    doc.add_paragraph(
        "Soft-delete workflow: Traveller DELETE and UPDATE commands set deletedAt timestamps and "
        "suppress records from active queries while preserving auditability — supporting PDPA-aligned "
        "data governance without destructive hard purges."
    )

    doc.add_page_break()

    # PAGE 6 — Issues Matrix
    add_heading(doc, "PAGE 6 — TECHNICAL ISSUES RESOLUTION MATRIX", 1)
    add_table(
        doc,
        ["Issue", "Technical Risk", "Resolution Architecture", "Verification"],
        [
            ("OTA margin compression", "Commercial", "5% commission engine in booking logic", "Transaction log audit"),
            ("Fragmented trip planning", "UX", "Unified planner + marketplace shell", "End-to-end UAT"),
            ("AI hallucinated venues", "Data quality", "Strict JSON schema prompts", "Schema validator"),
            ("Missing geocoordinates", "Map accuracy", "Google Places enrichment", "≥80% spot-check"),
            ("Zigzag map routes", "UX", "Nearest-neighbour + 2-opt sorting", "Path-length comparison"),
            ("Itinerary hard-delete loss", "Privacy/integrity", "Soft-delete + chat coordinator", "API state tests"),
            ("Unverified listings", "Trust", "Admin PENDING→APPROVED workflow", "Admin demo"),
            ("Single AI provider outage", "Availability", "Groq→Gemini→OpenAI fallback", "Failover test"),
            ("Payment misconfiguration", "Financial", "Stripe sandbox + webhook validation", "Receipt log"),
            ("Credential exposure", "Security", ".env gitignore; JWT; bcrypt", "Secret scan"),
        ],
    )

    doc.add_page_break()

    # PAGE 7 — Management
    add_heading(doc, "PAGE 7 — MANAGEMENT PLAN, CONDUCT & ROLES", 1)
    add_heading(doc, "G.1 Team Code of Conduct", 2)
    for item in [
        "Respect and inclusion across Business and IT members",
        "Transparency: blockers reported within 24 hours via GitHub Issues",
        "Accountability: sprint acceptance criteria must pass before merge",
        "Quality first: peer review on critical paths",
        "Academic integrity: all sources cited",
        "Security: no API keys committed to version control",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "G.2 Tiered Decision-Making Framework", 2)
    add_table(
        doc,
        ["Decision Class", "Authority", "Process"],
        [
            ("Day-to-day implementation", "Role owner", "Execute within sprint scope"),
            ("Cross-functional integration", "Weekly sync", "Simple majority vote"),
            ("Scope change (MVP features)", "Full squad + PM", "Written change log"),
            ("Deadlock", "Project Manager", "Binding resolution in minutes"),
        ],
    )

    add_heading(doc, "G.3 Workflow & File Management", 2)
    for item in [
        "Version control: GitHub (main ← develop ← feature/*); mandatory PR review",
        "Tracking: Trello/Jira board mapped to 10-week Gantt",
        "Documentation: Google Drive + /docs repository folder",
        "Communication: Team messaging; formal decisions in meeting minutes",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "G.4 Conflict Resolution", 2)
    for item in [
        "Direct resolution within 48 hours",
        "Weekly sync mediation",
        "PM binding decision",
        "Academic supervisor escalation if unresolved",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_page_break()

    # PAGE 8 — Gantt
    add_heading(doc, "PAGE 8 — ACTION PLAN & GANTT MILESTONES", 1)
    add_table(
        doc,
        ["Week", "Phase", "Key Deliverables", "Owner(s)", "Milestone", "Status"],
        [
            ("W1", "Initiation", "Charter, requirements, repo, literature", "PM + BA", "Kick-off", "Complete"),
            ("W2", "Design", "ERD, OpenAPI, wireframes, ADR", "UX + Backend", "Design sign-off", "Complete"),
            ("W3–4", "Core API", "Prisma schema, JWT, property modules", "Backend", "API alpha", "Complete"),
            ("W4–5", "Frontend shell", "Landing, search, auth, property detail", "Frontend", "UI shell demo", "Complete"),
            ("W5–6", "Payments", "Stripe sandbox, booking flow", "Backend", "Booking demo", "Complete"),
            ("W6–7", "Business & Admin", "Owner portal, admin approval", "Frontend + Backend", "Portal demo", "Complete"),
            ("W7–8", "AI Planner (parallel)", "LLM JSON, Places enrichment", "AI/DevOps", "AI alpha", "Complete"),
            ("W7–8", "Maps (parallel)", "Route optimisation, polylines", "Frontend", "Map demo", "Complete"),
            ("W9", "Integration", "AI chat, soft-delete CRUD, navigation", "All IT + PM", "Integration beta", "Complete"),
            ("W10", "Closure", "UAT, security, staging, docs", "PM + All", "Submission", "In progress"),
        ],
    )

    add_heading(doc, "Gantt Chart (Weeks 1–10)", 2)
    gantt_headers = ["Task / Phase", "W1", "W2", "W3", "W4", "W5", "W6", "W7", "W8", "W9", "W10"]
    gantt_rows = [
        ("Requirements & Setup", "██", "", "", "", "", "", "", "", "", ""),
        ("ERD, API, UI Design", "", "██", "", "", "", "", "", "", "", ""),
        ("Auth & Core API", "", "", "██", "██", "", "", "", "", "", ""),
        ("Frontend Shell", "", "", "", "██", "██", "", "", "", "", ""),
        ("Booking & Stripe", "", "", "", "", "██", "██", "", "", "", ""),
        ("Owner & Admin Portal", "", "", "", "", "", "██", "██", "", "", ""),
        ("AI Itinerary + Places", "", "", "", "", "", "", "██", "██", "", ""),
        ("Maps & Route UX", "", "", "", "", "", "", "██", "██", "", ""),
        ("Chat, CRUD, Soft-Delete", "", "", "", "", "", "", "", "", "██", ""),
        ("UAT, Deploy, Documentation", "", "", "", "", "", "", "", "", "", "██"),
    ]
    add_table(doc, gantt_headers, gantt_rows)

    doc.add_page_break()

    # PAGE 9 — Risk & References
    add_heading(doc, "PAGE 9 — RISK MANAGEMENT & ACADEMIC BIBLIOGRAPHY", 1)
    add_heading(doc, "H. Risk Management Matrix", 2)
    add_table(
        doc,
        ["Risk ID", "Risk", "L", "I", "Mitigation Strategy", "Residual"],
        [
            ("R1", "Low pilot listing count (<5)", "M", "M", "Seed demo properties; owner test accounts", "Low"),
            ("R2", "AI API quota / provider outage", "M", "H", "Groq→Gemini→OpenAI fallback; cache", "Low"),
            ("R3", "<80% geocoding accuracy", "M", "M", "Places validation; manual regeneration", "Low–M"),
            ("R4", "Scope creep (mobile, 50 hotels in MVP)", "H", "H", "Strict Part A/B partition", "Low"),
            ("R5", "Stripe integration delay", "M", "M", "Sandbox by W5; manual UAT fallback", "Low"),
            ("R6", "Maps API cost/limit", "M", "M", "Lazy map load; persist coords in DB", "Low"),
            ("R7", "Data breach / PDPA", "L", "C", "HTTPS, bcrypt, soft-delete, checklist", "Low"),
            ("R8", "Team member unavailability", "M", "M", "Cross-training; /docs knowledge base", "M"),
            ("R9", "Unequal contribution", "L–M", "M", "Sprint board ownership; PM escalation", "Low"),
            ("R10", "Unrealistic claims to assessors", "M", "H", "SMART KPIs; capstone vs beta split", "Low"),
        ],
    )

    add_heading(doc, "I. Capstone MVP Success Metrics", 2)
    add_table(
        doc,
        ["#", "Metric", "Target", "Evidence"],
        [
            ("1", "Staging deployment live", "100% core features", "UAT + URL"),
            ("2", "Pilot hotel listings", "≥5 verified", "DB count"),
            ("3", "5% commission engine", "5.00% on test bookings", "SQL audit"),
            ("4", "Booking completion rate", "≥95% (5 tests)", "Stripe logs"),
            ("5", "AI coordinate accuracy", "≥80% valid Places coords", "20-itinerary check"),
            ("6", "Soft-delete integrity", "100% (5 scenarios)", "API + deletedAt audit"),
            ("7", "Documentation complete", "100% inventory", "/docs checklist"),
            ("8", "Traveller UAT satisfaction", "≥80% rate ≥4/5 (n≥10)", "Likert survey"),
        ],
    )

    add_heading(doc, "J. References", 2)
    refs = [
        "Hospitality Net. (2024). The real cost of Booking.com: Five practical steps for Southeast Asian hotels. https://www.hospitalitynet.org/opinion/4132278/",
        "Hospitality Technology. (2024). Expedia and Booking.com increase their fees for independent hotels. https://hospitalitytech.com/expedia-and-booking-increase-their-fees-independent-hotels",
        "Juniarta, I. G. M., Wahyuni, C. C., Purnomo, P. D., & Palupiningtyas, D. (2026). More rooms, less profit: Unveiling OTA distribution channels' contradictory hotel impacts. Momentum Matrix: International Journal of Communication, Tourism, and Social Economic Trends, 3(1). https://doi.org/10.62951/momat.v3i1.603",
        "O'Connor, P., Piccoli, G., & Roth, M. (2025). Layered distribution penalties: Calculating the true net cost of OTA dependencies for independent economy properties. International Journal of Hospitality Management, 118, Article 103842.",
        "PhocusWire. (2024). Booking.com and Expedia commission hikes spark backlash from independent hotels. https://www.phocuswire.com/booking-expedia-commission-hikes-independent-hotels",
        "Rahadi, R. A., Handayati, R., & Utama, A. (2021). Assessing bottom-line profit margin compression within independent lodging structures across Southeast Asian emerging markets. Southeast Asian Journal of Business Studies, 10(2), 189–203.",
        "Zhong, L., & Moon, S. (2020). Distribution cost structures and dynamic channel mix optimizations for independent economy lodging providers. Tourism and Hospitality Research, 22(1), 74–88.",
    ]
    for i, ref in enumerate(refs, 1):
        doc.add_paragraph(f"{i}. {ref}")

    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    er = end.add_run("— End of Proposal — Version 3.0")
    er.bold = True
    er.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    return doc


def main():
    doc = build_document()
    OUT_DOCS.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCS)
    print(f"Saved: {OUT_DOCS}")

    try:
        doc.save(OUT_DOWNLOADS)
        print(f"Saved: {OUT_DOWNLOADS}")
    except PermissionError:
        alt = OUT_DOWNLOADS.with_name("BX3102_A1_Proposal_v3_NEW.docx")
        doc.save(alt)
        print(f"Downloads path locked; saved: {alt}")

    # PDF export attempt
    try:
        from docx2pdf import convert
        pdf_docs = OUT_DOCS.with_suffix(".pdf")
        convert(str(OUT_DOCS), str(pdf_docs))
        print(f"PDF: {pdf_docs}")
        pdf_dl = OUT_DOWNLOADS.with_suffix(".pdf")
        convert(str(OUT_DOCS), str(pdf_dl))
        print(f"PDF: {pdf_dl}")
    except Exception as e:
        print(f"PDF export skipped: {e}")


if __name__ == "__main__":
    main()
