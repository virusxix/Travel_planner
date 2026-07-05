"""Set Section D to exact stakeholder table only."""
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

DOC_PATH = Path(r"d:\sample\HiddenStay_v4_Final.docx")

ROWS = [
    ("Travellers", "Affordable stays + map-ready itineraries", "Web UX, search/book, AI with Places coords"),
    (
        "Independent Hotel Owners",
        "Lower commission, visibility",
        "5% fee, self-service dashboard, pilot notifications",
    ),
    (
        "Project Team & Academic Assessors",
        "On-time MVP, competence demo",
        "C.7 ownership, staging deploy, SMART evidence",
    ),
]


def delete_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def find_stakeholder_table(doc: Document):
    for table in doc.tables:
        if table.rows and table.rows[0].cells[0].text.strip() == "Stakeholder":
            return table
    return None


def main():
    doc = Document(DOC_PATH)

    # Remove intro paragraph(s) between D heading and F References
    d_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "D.   Stakeholder Analysis":
            d_idx = i
            break

    if d_idx is not None:
        to_remove = []
        for p in doc.paragraphs[d_idx + 1 :]:
            t = p.text.strip()
            if t.startswith("F.   References"):
                break
            if t:
                to_remove.append(p)
        for p in to_remove:
            delete_paragraph(p)

    table = find_stakeholder_table(doc)
    if table is None:
        raise RuntimeError("Stakeholder table not found")

    headers = ["Stakeholder", "Primary Interests", "Capstone MVP Expectations"]
    for ci, h in enumerate(headers):
        table.rows[0].cells[ci].text = h

    for ri, row_data in enumerate(ROWS, start=1):
        for ci, val in enumerate(row_data):
            table.rows[ri].cells[ci].text = val

    doc.save(DOC_PATH)
    print(f"Updated: {DOC_PATH}")


if __name__ == "__main__":
    main()
