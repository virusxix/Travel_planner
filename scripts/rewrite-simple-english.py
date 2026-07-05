"""Rewrite HiddenStay_v4_Final.docx in simpler undergraduate English."""
from pathlib import Path

from docx import Document

DOC_PATH = Path(r"d:\sample\HiddenStay_v4_Final.docx")

# Paragraph replacements: exact match or startswith key -> new text
PARA_REPLACEMENTS = {
    "Independent and small hotels across Southeast Asia": (
        "Many small and independent hotels in Southeast Asia rely heavily on big booking websites "
        "like Booking.com and Expedia to get customers. This helps them get bookings, but it also "
        "costs them a lot of money."
    ),
    "Commission fees charged by dominant OTA platforms": (
        "These platforms usually charge around 15% to 30% commission per booking. "
        "Juniarta et al. (2026) found that OTAs can increase room sales (β₁ = 0.897, p < 0.01), "
        "but they also reduce hotel profit margins (β₂ = −0.652, p < 0.01). "
        "Budget hotels are hit hardest, losing about SGD 3.47 EBITDA per available room on average."
    ),
    "Independent hotels in SEA route approximately": (
        "In Southeast Asia, independent hotels send about 61% of online bookings through OTAs, "
        "while chain hotels only use OTAs for about 35% (Hospitalitynet.org, 2024)."
    ),
    "Both Expedia and Booking.com have recently increased": (
        "Booking.com and Expedia have also raised fees for independent hotels in recent years "
        "(Hospitality Technology, 2024; PhocusWire, 2024)."
    ),
    "Economy and budget hotels experience net negative": (
        "Budget hotels often lose money after paying OTA fees. Luxury hotels can handle the cost better "
        "(Juniarta et al., 2026)."
    ),
    "A secondary problem affects the traveller side": (
        "Travellers face a second problem: trip planning is spread across many apps. "
        "Most AI travel tools give generic suggestions without real map locations or links to "
        "bookable small hotels."
    ),
    "HiddenStay addresses both problems": (
        "HiddenStay tries to solve both problems with one platform: a lower-commission booking site "
        "(5% base fee) and an AI trip planner that creates map-ready itineraries with real place data."
    ),
    "All objectives are formulated according to SMART": (
        "Our goals follow SMART criteria (Specific, Measurable, Achievable, Relevant, Time-bound)."
    ),
    "The Southeast Asian hotel industry is dominated": (
        "Most hotels in Southeast Asia are small independent businesses that cannot negotiate "
        "lower OTA fees like big hotel chains. The OTA market is large and growing, but there is "
        "still no strong platform built mainly for small independent hotels with fair commission "
        "and built-in trip planning."
    ),
    "At 5% commission, a hotel charging SGD 100": (
        "With a 5% fee, a SGD 100/night booking gives the hotel SGD 95. "
        "With a typical 20–30% OTA fee, the hotel only keeps SGD 70–80. "
        "For hotels with tight margins, this difference really matters."
    ),
    "Southeast Asia’s travel market has recovered": (
        "Travel in Southeast Asia has recovered well after COVID. More travellers want local and "
        "boutique stays, but small hotels are still hard to find without using expensive OTAs. "
        "HiddenStay gives them a cheaper way to get online bookings."
    ),
    "HiddenStay is developed by a cross-functional team": (
        "HiddenStay is built by a 5-person student team with both business and IT skills. "
        "The table below shows each member’s role and why they were chosen for that role."
    ),
    "HiddenStay is a Southeast Asia-focused OTA platform": (
        "HiddenStay is a booking platform focused on Southeast Asia and built for small independent hotels. "
        "Users can search, list, book, and pay through a website (and planned mobile-friendly design), "
        "plus use an AI trip planner in the same system."
    ),
    "The project is divided into two scopes": (
        "We split the project into two parts: (1) Capstone MVP — what we must finish in 10 weeks, "
        "and (2) Future Business Vision — what the business could become after graduation. "
        "This keeps our submission realistic for a student team."
    ),
    "The following features constitute the scope": (
        "These are the features we plan to build and demo during the 10-week capstone:"
    ),
    "Features explicitly excluded from Capstone MVP": (
        "We will NOT build these in the capstone (to avoid scope creep):"
    ),
    "The following represent the longer-term commercial roadmap": (
        "These are post-capstone business goals only (not graded in the 10-week MVP):"
    ),
    "Upon successful delivery, HiddenStay is expected": (
        "If the project goes well, these are the results we expect, split by capstone vs post-capstone:"
    ),
    "All HiddenStay team members commit to the following professional principles": (
        "Our team agrees to follow these basic rules:"
    ),
    "Respect and Inclusion:": "Respect: Treat every team member fairly and professionally.",
    "Transparency:": "Be honest: Share progress, problems, and risks early in weekly meetings.",
    "Accountability:": "Own your tasks: Finish your assigned work on time.",
    "Quality First:": "Check quality: Do not submit work that fails our sprint checklist.",
    "Punctuality:": "Be on time: Attend meetings. Tell the team early if you cannot attend.",
    "Constructive Feedback:": "Give useful feedback: Criticise the work, not the person.",
    "A full breakdown of team roles": (
        "Full role details are in Section A. Each person has clear tasks so work does not overlap. "
        "The Project Manager keeps the schedule and communication on track."
    ),
    "Decisions are made using the following tiered approach": (
        "We use simple decision rules:"
    ),
    "Day-to-day decisions:": "Small decisions: Each member decides within their own task area.",
    "Cross-functional decisions:": "Shared decisions: Discuss in weekly meeting and vote if needed.",
    "Major strategic decisions": (
        "Big decisions (scope change, major delays): All 5 members must agree and record it in meeting notes."
    ),
    "Deadlocked decisions:": "If we still cannot agree, the Project Manager makes the final call.",
    "Version control:": (
        "Code: Private GitHub repo (github.com/oakkar-cm/Travel), feature branches, pull request reviews."
    ),
    "Project tracking:": "Tasks: Tracked on Trello or Jira with weekly updates.",
    "Documentation:": "Files: Stored in Google Drive/OneDrive with clear version names.",
    "Communication:": "Chat: Team group chat for daily talk; meeting minutes for important decisions.",
    "Should a conflict arise between team members": (
        "If team members disagree, we follow these steps:"
    ),
    "The parties involved attempt to resolve": "Step 1: Talk directly and try to fix it within 48 hours.",
    "If unresolved, the issue is raised at the next team sync": "Step 2: Bring it to the weekly team meeting.",
    "If still unresolved, the Project Manager issues": "Step 3: Project Manager makes a final decision.",
    "All conflict resolutions are documented": "Step 4: Write the outcome in the project log.",
    "All work produced for HiddenStay must be original": (
        "All work must be original or properly cited. No plagiarism."
    ),
    "No team member shall misrepresent progress": (
        "Do not fake progress or hide blockers."
    ),
    "All external hotel partner communications must accurately": (
        "When talking to hotel partners, only promise features that actually work in the current build."
    ),
    "User data and hotel partner data will be handled": (
        "We will follow PDPA rules when handling user and partner data."
    ),
    "Note on Phase 6 workload distribution": (
        "Weeks 7–8 are heavy for AI work, so we split tasks: Lead Developer builds the AI engine and "
        "Google Places integration, while BA and PM work on prompts, JSON format, and itinerary UI pages."
    ),
    "The chart below illustrates the 10-week": (
        "The Gantt chart shows our 10-week plan. Different colours show team tracks running in parallel, "
        "especially AI work in Weeks 7–9."
    ),
    "The following risk register identifies": (
        "This table lists main project risks, how likely they are, and how we plan to handle them."
    ),
    "The following sources substantiate": (
        "These references support our problem statement and market background:"
    ),
    "Competitive Positioning (OTA Comparison)": "How HiddenStay Compares to Major OTAs",
    "Technical approach (AI module):": (
        "How the AI module works: We use structured JSON prompts, Google Places for real coordinates, "
        "and backup AI providers (Groq, Gemini, OpenAI) so the demo still works if one API fails."
    ),
    "AI module priority and contingency:": (
        "AI build order: (1) itinerary JSON + Places geocoding, (2) map view, (3) AI chat. "
        "If we run out of time in Weeks 7–9, we keep (1) and (2) and delay chat if needed."
    ),
}

# Table replacements by header signature -> {row_col or row index: text}
TABLE_DATA = {
    "objectives": {
        "headers": ["#", "Objective", "Success Indicator / KPI", "Measurement Method"],
        "rows": [
            ("1", "Build a working booking website for small hotels (search, list, book).", "Core booking flow works end-to-end on staging.", "UAT test checklist with ≥90% pass rate."),
            ("2", "Code a 5% commission fee into sandbox bookings.", "All test bookings store 5.00% commission.", "Check booking records in database + unit tests."),
            ("3", "Give hotel partners a simple dashboard (calendar, bookings, pricing, basic stats).", "4 dashboard features work; pilot partners rate it ≥4/5.", "Feature checklist + short partner survey."),
            ("4", "Add 5 real pilot hotel listings for testing.", "5 listings live with photos, prices, and availability.", "Count records in database + manual review."),
            ("5", "Build an AI trip planner with map-ready stops.", "≥80% of AI stops have valid Google Places coordinates.", "Run validation script on 30 sample itineraries."),
        ],
    },
    "issues": {
        "headers": ["Issue", "How HiddenStay Addresses It"],
        "rows": [
            ("High OTA fees", "HiddenStay charges 5% instead of 15–30%, so hotels keep more money."),
            ("No fair booking option for small hotels", "HiddenStay is built for independent SEA hotels, not big chains."),
            ("Weak hotel tools", "Hotel dashboard helps partners manage bookings, calendar, and pricing."),
            ("Hard to discover small hotels", "Search and filters make independent hotels easier to find."),
            ("Trip planning is scattered", "AI planner builds day-by-day plans in the same app used for booking."),
            ("Generic AI travel suggestions", "Google Places adds real venue names and map coordinates."),
        ],
    },
    "competitor": {
        "headers": ["Feature", "Booking.com", "Expedia", "HiddenStay"],
        "rows": [
            ("Commission", "15–30%", "15–30%", "5% (capstone MVP)"),
            ("Focus on independent hotels", "Partial", "Partial", "Yes"),
            ("Built-in AI trip planner", "No", "Limited", "Yes"),
            ("Map-based itineraries", "No", "No", "Yes"),
        ],
    },
    "outcomes": {
        "headers": ["Outcome", "Target", "Phase / Timeline"],
        "rows": [
            ("Staging site live", "Core MVP features work on a public test URL", "Capstone — Week 10"),
            ("Pilot hotel listings", "5 test hotels with photos, prices, availability", "Capstone — Week 10"),
            ("Booking flow works", "≥90% of UAT test bookings complete without error", "Capstone — Week 10"),
            ("AI map accuracy", "≥80% of AI stops have valid Google Places data", "Capstone — Week 10"),
            ("More hotel partners", "50 partners in Singapore and Malaysia", "After capstone — Month 9"),
            ("Monthly bookings", "500+ bookings/month", "After capstone — Month 12"),
            ("Large partner base", "500+ hotels across 6 SEA countries", "After capstone — Month 24"),
            ("Break-even", "Revenue covers operating costs", "After capstone — Month 24"),
        ],
    },
    "stakeholder": {
        "headers": ["Stakeholder", "Interest", "Influence", "Capstone MVP Expectations"],
        "rows": [
            ("Travellers", "High", "Medium", "Easy booking, fair prices, AI trips with real map locations."),
            ("Independent Hotel Owners", "High", "Medium", "5% fee, listing tools, booking alerts during pilot."),
            ("Project Team & Lecturers", "High", "High", "Working demo on time, clear docs, proof of SMART goals."),
            ("API Providers (Google, Stripe, AI)", "Medium", "Low", "Sandbox APIs work reliably; fallback if one service fails."),
        ],
    },
    "org": {
        "headers": ["No.", "Member Name", "Role", "Key Responsibilities", "Skills / Role Justification (Coursework-Linked)"],
        # Keep names from doc - only simplify justification column
    },
    "risks": {
        "rows": [
            ("Low hotel partner adoption during pilot phase", "Medium", "High",
             "Offer free onboarding help and test accounts; use team contacts for first listings.",
             "Low–Medium"),
            ("Lead Developer overload in Weeks 7–8", "High", "High",
             "Split AI work: developer handles APIs; BA/PM handles prompts and UI.",
             "Low"),
            ("Project runs over time", "Medium", "Medium",
             "Use weekly sprints and cut non-essential features before missing core demo.",
             "Low"),
            ("Stripe integration delays", "Medium", "Medium",
             "Start Stripe sandbox early; keep manual test booking as backup for demo.",
             "Low"),
            ("AI API limits or downtime", "Medium", "Medium",
             "Use Groq → Gemini → OpenAI fallback and cache old itineraries.",
             "Low"),
            ("Data/security issues", "Low", "Very High",
             "Use HTTPS, hashed passwords, and soft-delete for user data (PDPA).",
             "Low"),
            ("Team member unavailable", "Low", "High",
             "Share docs in GitHub/Drive and cross-train on key modules.",
             "Medium"),
        ],
    },
}

BULLET_SIMPLIFY = {
    "User authentication (JWT-based": "User login/sign-up for travellers and hotel partners (JWT auth)",
    "Hotel listing module": "Hotel partners can add listings with photos, rooms, prices, and availability",
    "Search and filter functionality": "Search hotels by destination, dates, price, and property type",
    "Core booking engine": "Book rooms and pay with Stripe sandbox (test mode only, no real money)",
    "Hotel partner dashboard": "Hotel dashboard for bookings, calendar, and basic stats",
    "AI Trip Planner module": "AI trip planner using Groq/Gemini/OpenAI + Google Places + map view",
    "Traveller dashboard": "Traveller dashboard for saved trips, bookings, and AI chat",
    "Review and rating system": "Guest reviews after stays",
    "Deployment to a publicly": "Deploy to a live staging URL (Vercel + Supabase)",
    "Live payment processing": "Real payment processing (we only use Stripe test mode)",
    "Multi-language localisation": "Multiple languages",
    "Mobile app store submission": "Publishing native iOS/Android apps",
    "Full SEA market expansion": "Expansion to all SEA countries (capstone focuses on SG/MY testing)",
    "Advanced revenue management": "Advanced dynamic pricing engine",
    "Full public launch": "Full launch in 6 SEA countries",
    "200+ hotel partners": "200+ hotel partners and 500+ bookings/month by Month 12",
    "Promotional listing tiers": "Higher 10–15% promo tiers with ads",
    "Native mobile app release": "Release iOS and Android apps",
    "Platform break-even within": "Break-even within 24 months",
    "Investor pitch and Series A": "Investor pitch and funding preparation",
}


def replace_paragraph_text(paragraph, new_text: str):
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def apply_para_replacements(doc: Document):
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        for key, new in PARA_REPLACEMENTS.items():
            if t.startswith(key) or t == key:
                replace_paragraph_text(p, new)
                break
        else:
            for key, new in BULLET_SIMPLIFY.items():
                if t.startswith(key):
                    replace_paragraph_text(p, new)
                    break


def identify_table(table):
    if not table.rows:
        return None
    h = [c.text.strip() for c in table.rows[0].cells]
    joined = " ".join(h)
    if h[0] == "#":
        return "objectives"
    if h[0] == "Issue":
        return "issues"
    if h[0] == "Feature" and "HiddenStay" in joined:
        return "competitor"
    if h[0] == "Outcome":
        return "outcomes"
    if h[0] == "Stakeholder" and "Influence" in joined:
        return "stakeholder"
    if h[0] == "Risk":
        return "risks"
    if h[0] == "No." and "Skills" in joined:
        return "org"
    return None


def simplify_org_justification(text: str) -> str:
    mapping = {
        "Proven execution": "Completed project management and Agile coursework; organised past team projects.",
        "Advanced engineering": "Software engineering and web development modules; built full-stack apps before.",
        "Practical translation": "HCI and UI/UX coursework; experience with Figma and user testing.",
        "Background in Business": "Business analysis and economics modules; comfortable with numbers and market research.",
        "Knowledge of Digital": "Digital marketing coursework; helped with outreach and presentation tasks.",
    }
    for k, v in mapping.items():
        if k in text:
            return v
    if len(text) > 120:
        return text[:120].rsplit(" ", 1)[0] + "."
    return text


def apply_table_replacements(doc: Document):
    for table in doc.tables:
        kind = identify_table(table)
        if kind == "org":
            for ri in range(1, len(table.rows)):
                cell = table.rows[ri].cells[4]
                cell.text = simplify_org_justification(cell.text.strip())
            continue
        if kind not in TABLE_DATA:
            continue
        data = TABLE_DATA[kind]
        if "headers" in data:
            for ci, h in enumerate(data["headers"]):
                if ci < len(table.rows[0].cells):
                    table.rows[0].cells[ci].text = h
        rows = data["rows"]
        while len(table.rows) > len(rows) + 1:
            table._tbl.remove(table.rows[-1]._tr)
        while len(table.rows) < len(rows) + 1:
            table.add_row()
        for ri, row_vals in enumerate(rows, start=1):
            for ci, val in enumerate(row_vals):
                if ci < len(table.rows[ri].cells):
                    table.rows[ri].cells[ci].text = val


def simplify_juniarta_box(doc: Document):
    for table in doc.tables:
        if table.rows and "Juniarta" in table.rows[0].cells[0].text:
            if len(table.rows) >= 4:
                table.rows[0].cells[0].text = "Key finding — Juniarta et al. (2026)"
                table.rows[1].cells[0].text = (
                    '"More Rooms, Less Profit" — OTAs increase bookings but reduce profit.'
                )
                table.rows[2].cells[0].text = "β₁ = 0.897: more OTA use → more room sales."
                table.rows[3].cells[0].text = "β₂ = −0.652: more OTA use → lower profit margins."
                if len(table.rows) > 4:
                    table.rows[4].cells[0].text = (
                        "Budget hotels: about −SGD 3.47 per room. 90.3% of managers say OTA fees are too high."
                    )


def main():
    doc = Document(DOC_PATH)
    apply_para_replacements(doc)
    apply_table_replacements(doc)
    simplify_juniarta_box(doc)

    try:
        doc.save(DOC_PATH)
        saved = DOC_PATH
    except PermissionError:
        saved = DOC_PATH.with_name("HiddenStay_v4_Final_Simple.docx")
        doc.save(saved)

    print(f"Saved: {saved}")


if __name__ == "__main__":
    main()
