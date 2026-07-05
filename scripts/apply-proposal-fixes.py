"""Apply assessor feedback fixes to HiddenStay_v4_Final.docx."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

DOC_PATH = Path(r"d:\sample\HiddenStay_v4_Final.docx")

# --- Objectives: capstone-only (remove old #5 break-even, #7; renumber) ---
OBJECTIVES = [
    (
        "1",
        "Build a functional web OTA enabling independent hotel listing, search, and booking within the 10-week capstone.",
        "Functional staging platform with core booking flow operational end-to-end.",
        "UAT completion with ≥90% task-completion rate across defined user journeys.",
    ),
    (
        "2",
        "Implement a 5% base commission model in the booking payment engine (sandbox).",
        "Commission field persisted at 5.00% on all test bookings.",
        "SQL audit of booking transaction logs; unit test assertions on fee computation.",
    ),
    (
        "3",
        "Empower hotel partners with a self-service dashboard (calendar, bookings, pricing, basic analytics).",
        "≥4 core dashboard features functional; ≥80% of pilot partners rate usefulness ≥4/5.",
        "Feature checklist; post-onboarding partner survey (Likert scale).",
    ),
    (
        "4",
        "Onboard 5 verified pilot hotel listings for capstone system stability testing.",
        "5 live listings with photos, room types, pricing, and availability.",
        "Database property count at Week 10; manual listing quality review.",
    ),
    (
        "5",
        "Deliver an AI Trip Planner generating multi-day, geocoded, map-ready itineraries.",
        "≥80% of AI-generated stops contain valid Google Places coordinates.",
        "Automated coordinate validation script against 30 sample itineraries; QA report.",
    ),
]

AI_CONTINGENCY = (
    "AI module priority and contingency: (1) structured JSON itinerary generation and Google Places "
    "geocoding, (2) map visualisation with route rendering, (3) conversational AI chat refinement. "
    "If API integration risks arise during Weeks 7–9, the MVP will prioritise itinerary generation "
    "and map visualisation; advanced AI chat refinement may be deferred without blocking capstone "
    "submission."
)

COMPETITOR_HEADERS = ["Feature", "Booking.com", "Expedia", "HiddenStay"]
COMPETITOR_ROWS = [
    ("Commission (independent hotels)", "15–30%", "15–30%", "5% base (capstone MVP)"),
    ("Independent hotel focus", "Partial", "Partial", "Yes — core target segment"),
    ("Integrated AI trip planner", "No", "Limited", "Yes — structured JSON + Places"),
    ("Map-ready geocoded itineraries", "No", "No", "Yes — Google Maps integration"),
]

STAKEHOLDER_HEADERS = ["Stakeholder", "Interest", "Influence", "Capstone MVP Expectations"]
STAKEHOLDER_ROWS = [
    (
        "Travellers",
        "High",
        "Medium",
        "Affordable stays + map-ready itineraries; web UX, search/book, AI with Places coords.",
    ),
    (
        "Independent Hotel Owners",
        "High",
        "Medium",
        "Lower commission, visibility; 5% fee, self-service dashboard, pilot notifications.",
    ),
    (
        "Project Team & Academic Assessors",
        "High",
        "High",
        "On-time MVP, competence demo; C.7 ownership, staging deploy, SMART evidence.",
    ),
    (
        "API & Payment Providers (Google, Stripe, LLM)",
        "Medium",
        "Low",
        "Reliable sandbox integrations; quota/rate limits managed via fallback and caching (C.9).",
    ),
]

HEADER_FILL = "1A365D"
LABEL_FILL = "EDF2F7"


def set_cell_shading(cell, fill: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.find(qn("w:shd"))
    if old is not None:
        tc_pr.remove(old)
    tc_pr.append(shading)


def write_cell(cell, text: str, *, header=False, label=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(9 if not header else 10)
    run.font.bold = header or label
    if header:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, HEADER_FILL)
    elif label:
        set_cell_shading(cell, LABEL_FILL)
    else:
        set_cell_shading(cell, "FFFFFF")


def fill_table(table, headers, rows, label_col=0):
    for ci, h in enumerate(headers):
        write_cell(table.rows[0].cells[ci], h, header=True)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            write_cell(table.rows[ri].cells[ci], val, label=(ci == label_col))


def find_table_by_header(doc, first_header: str):
    for table in doc.tables:
        if table.rows and table.rows[0].cells[0].text.strip() == first_header:
            return table
    return None


def insert_after(element, new_element):
    element.addnext(new_element)


def add_paragraph_after(paragraph, text: str, bold=False):
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    if bold:
        rpr = OxmlElement("w:rPr")
        rpr.append(OxmlElement("w:b"))
        r.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    insert_after(paragraph._element, p)
    return p


def rebuild_objectives_table(doc):
    table = find_table_by_header(doc, "#")
    if table is None:
        raise RuntimeError("Objectives table not found")

    # Resize: header + 5 rows
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    while len(table.rows) < 6:
        table.add_row()

    fill_table(
        table,
        ["#", "Objective", "Success Indicator / KPI", "Measurement Method"],
        OBJECTIVES,
        label_col=0,
    )


def add_competitor_table(doc):
    if any(
        table.rows
        and table.rows[0].cells[0].text.strip() == "Feature"
        and len(table.rows[0].cells) >= 4
        and "HiddenStay" in table.rows[0].cells[3].text
        for table in doc.tables
    ):
        return

    anchor = None
    for p in doc.paragraphs:
        if p.text.strip() == "B.2   Issues to Be Addressed":
            anchor = p
            break

    issues_table = find_table_by_header(doc, "Issue")
    insert_point = issues_table._element if issues_table is not None else anchor._element

    heading = OxmlElement("w:p")
    hr = OxmlElement("w:r")
    hrpr = OxmlElement("w:rPr")
    hrpr.append(OxmlElement("w:b"))
    hr.append(hrpr)
    ht = OxmlElement("w:t")
    ht.text = "Competitive Positioning (OTA Comparison)"
    hr.append(ht)
    heading.append(hr)
    insert_after(insert_point, heading)

    new_table = doc.add_table(rows=1 + len(COMPETITOR_ROWS), cols=4)
    fill_table(new_table, COMPETITOR_HEADERS, COMPETITOR_ROWS, label_col=0)
    insert_after(heading, new_table._element)


def add_ai_contingency(doc):
    if any(AI_CONTINGENCY[:40] in p.text for p in doc.paragraphs):
        return

    anchor = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("Technical approach (AI module):"):
            anchor = p
            break
    if anchor is None:
        for p in doc.paragraphs:
            if p.text.strip().startswith("Features explicitly excluded from Capstone MVP"):
                anchor = doc.paragraphs[max(0, doc.paragraphs.index(p) - 1)]
                break

    if anchor:
        add_paragraph_after(anchor, AI_CONTINGENCY)


def rebuild_stakeholder_table(doc):
    table = find_table_by_header(doc, "Stakeholder")
    if table is None:
        raise RuntimeError("Stakeholder table not found")

    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    while len(table.rows) < 5:
        table.add_row()

    # Ensure 4 columns
    # python-docx tables have fixed cols from creation - check cols
    if len(table.columns) < 4:
        # rebuild by delete + add
        parent = table._element.getparent()
        idx = list(parent).index(table._element)
        parent.remove(table._element)
        new_table = doc.add_table(rows=1 + len(STAKEHOLDER_ROWS), cols=4)
        fill_table(new_table, STAKEHOLDER_HEADERS, STAKEHOLDER_ROWS, label_col=0)
        parent.insert(idx, new_table._element)
    else:
        fill_table(table, STAKEHOLDER_HEADERS, STAKEHOLDER_ROWS, label_col=0)


def update_b4_future_vision(doc):
    """Ensure moved commercial goals are explicit in B.4."""
    extra_bullets = [
        "≥500 completed bookings per month by end of Year 1 (commercial traction target)",
        "Platform break-even within 24 months of full launch (commission revenue ≥ operating costs)",
    ]
    in_b4 = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "B.4   Future Business Vision":
            in_b4 = True
            continue
        if in_b4 and t.startswith("B.5"):
            break
        if in_b4:
            for bullet in extra_bullets:
                if bullet.split("(")[0].strip()[:20] in t:
                    return  # already present

    # Insert before B.5
    for p in doc.paragraphs:
        if p.text.strip() == "B.5   Expected Project Outcomes":
            for bullet in reversed(extra_bullets):
                add_paragraph_after(doc.paragraphs[doc.paragraphs.index(p) - 1], f"• {bullet}")
            break


def main():
    doc = Document(DOC_PATH)
    rebuild_objectives_table(doc)
    add_competitor_table(doc)
    add_ai_contingency(doc)
    rebuild_stakeholder_table(doc)
    update_b4_future_vision(doc)

    try:
        doc.save(DOC_PATH)
        saved = DOC_PATH
    except PermissionError:
        saved = DOC_PATH.with_name("HiddenStay_v4_Final_Revised.docx")
        doc.save(saved)

    print(f"Saved: {saved}")
    print("Applied: objectives trim, competitor table, AI contingency, stakeholder matrix, B.4 commercial goals")


if __name__ == "__main__":
    main()
