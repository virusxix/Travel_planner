"""Trim Section D stakeholder table in HiddenStay_v4_Final.docx."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

DOC_PATH = Path(r"d:\sample\HiddenStay_v4_Final.docx")

INTRO = (
    "A stakeholder analysis identifies the groups with a material interest in HiddenStay "
    "and their expectations for the capstone MVP. Post-capstone commercial stakeholders "
    "(e.g. investors and regional partners) are addressed in Section B.4 only."
)

HEADERS = ["Stakeholder", "Primary Interests", "Capstone MVP Expectations"]

ROWS = [
    (
        "Travellers",
        "Discover and book authentic, affordable independent accommodation; access personalised, "
        "map-ready trip itineraries in a single platform.",
        "Intuitive web UX; search and filter; reliable booking confirmation; AI itineraries with "
        "valid Google Places coordinates.",
    ),
    (
        "Independent Hotel Owners",
        "Reduce OTA commission burden; increase direct online visibility; manage bookings without "
        "prohibitive platform fees.",
        "Transparent 5% base commission; self-service listing and dashboard; booking notifications "
        "during the 5-hotel pilot phase.",
    ),
    (
        "Project Team & Academic Assessors",
        "Deliver a functional MVP within the 10-week capstone timeline; demonstrate interdisciplinary "
        "business and technical competence.",
        "Clear task ownership per C.7; working staging deployment; SMART objective evidence "
        "(see Objectives and B.5).",
    ),
]

HEADER_FILL = "1A365D"
LABEL_FILL = "EDF2F7"


def set_cell_shading(cell, fill: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.find(qn("w:shd"))
    if old is not None:
        tc_pr.remove(old)
    tc_pr.append(shading)


def set_cell_text(cell, text: str, *, bold=False, header=False, label=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(9 if not header else 10)
    run.font.bold = bold or header
    if header:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, HEADER_FILL)
    elif label:
        run.font.bold = True
        set_cell_shading(cell, LABEL_FILL)
    else:
        set_cell_shading(cell, "FFFFFF")


def find_stakeholder_table(doc: Document):
    for table in doc.tables:
        if not table.rows:
            continue
        hdr = [c.text.strip() for c in table.rows[0].cells]
        if hdr and hdr[0] == "Stakeholder" and "Engagement" in " ".join(hdr):
            return table
    return None


def delete_table(table):
    table._element.getparent().remove(table._element)


def build_stakeholder_table(doc: Document):
    table = doc.add_table(rows=1 + len(ROWS), cols=3)
    for ci, h in enumerate(HEADERS):
        set_cell_text(table.rows[0].cells[ci], h, header=True)

    for ri, (stakeholder, interests, expectations) in enumerate(ROWS, start=1):
        row = table.rows[ri]
        set_cell_text(row.cells[0], stakeholder, label=True)
        set_cell_text(row.cells[1], interests)
        set_cell_text(row.cells[2], expectations)

    return table


def insert_after_paragraph(paragraph, element):
    paragraph._element.addnext(element)


def main():
    doc = Document(DOC_PATH)

    d_heading = None
    d_intro = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "D.   Stakeholder Analysis":
            d_heading = p
        if d_heading and d_intro is None and t.startswith("A stakeholder analysis"):
            d_intro = p

    if d_intro:
        d_intro.clear()
        run = d_intro.add_run(INTRO)
        run.font.size = Pt(11)

    old_table = find_stakeholder_table(doc)
    anchor = d_intro or d_heading
    if old_table is not None:
        delete_table(old_table)

    if anchor is None:
        raise RuntimeError("Section D not found")

    new_table = build_stakeholder_table(doc)
    insert_after_paragraph(anchor, new_table._element)

    try:
        doc.save(DOC_PATH)
        saved = DOC_PATH
    except PermissionError:
        saved = DOC_PATH.with_name("HiddenStay_v4_Final_TrimmedD.docx")
        doc.save(saved)

    print(f"Saved: {saved}")
    print("Section D: 3 stakeholders x 3 columns (removed Engagement Strategy, Admin, Investors)")


if __name__ == "__main__":
    main()
