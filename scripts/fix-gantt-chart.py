"""Fix Gantt chart table in HiddenStay_v4_Final.docx — shaded cells + legend."""
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

DOC_PATH = Path(r"d:\sample\HiddenStay_v4_Final.docx")

NAVY = "1A365D"
TEAL = "0D9488"
PURPLE = "6B46C1"
WHITE = "FFFFFF"
LIGHT_GRAY = "F7FAFC"

# row index -> list of (week_col_index, color)  [col 0 = label, cols 1-10 = W1-W10]
GANTT_ROWS = [
    ("1. Initiation", [(1, NAVY), (2, NAVY)]),
    ("2. Requirements", [(3, NAVY), (4, NAVY)]),
    ("3. Design", [(5, NAVY), (6, NAVY)]),
    ("4. Dev Sprint 1 (Lead Dev)", [(7, TEAL), (8, TEAL)]),
    ("5. Dev Sprint 2 (Lead Dev)", [(9, TEAL), (10, TEAL)]),
    ("6. AI Planner — Dev track", [(7, TEAL), (8, TEAL)]),
    ("6. AI Planner — BA/PM support", [(7, PURPLE), (8, PURPLE), (9, PURPLE)]),
    ("7. Testing & QA", [(9, NAVY)]),
    ("8. Deploy & Review", [(10, NAVY)]),
]


def set_cell_shading(cell, fill: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.find(qn("w:shd"))
    if old is not None:
        tc_pr.remove(old)
    tc_pr.append(shading)


def clear_cell(cell):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def style_header_cell(cell, text: str):
    clear_cell(cell)
    cell.text = text
    set_cell_shading(cell, NAVY)
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)


def style_label_cell(cell, text: str):
    clear_cell(cell)
    cell.text = text
    set_cell_shading(cell, LIGHT_GRAY)
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = False
            run.font.size = Pt(9)


def style_week_cell(cell, color: str | None):
    clear_cell(cell)
    if color:
        set_cell_shading(cell, color)
    else:
        set_cell_shading(cell, WHITE)


def find_gantt_table(doc: Document):
    for table in doc.tables:
        if not table.rows:
            continue
        header = [c.text.strip() for c in table.rows[0].cells]
        if header[:3] == ["Phase / Track", "W1", "W2"]:
            return table
    raise RuntimeError("Gantt table not found")


def find_legend_table(doc: Document, gantt_table):
    gantt_el = gantt_table._tbl
    for table in doc.tables:
        if table._tbl is gantt_el:
            continue
        if len(table.rows) == 1 and len(table.columns) >= 3:
            texts = " ".join(c.text for c in table.rows[0].cells)
            if "Lead Dev" in texts or "All-team" in texts:
                return table
    return None


def fix_gantt_table(table):
    headers = ["Phase / Track", "W1", "W2", "W3", "W4", "W5", "W6", "W7", "W8", "W9", "W10"]
    for ci, h in enumerate(headers):
        style_header_cell(table.rows[0].cells[ci], h)

    for ri, (label, fills) in enumerate(GANTT_ROWS, start=1):
        row = table.rows[ri]
        style_label_cell(row.cells[0], label)
        fill_map = dict(fills)
        for ci in range(1, 11):
            style_week_cell(row.cells[ci], fill_map.get(ci))


def fix_legend_table(table):
    """Rebuild legend as a compact 1-row color key (handles merged-cell duplicates)."""
    row = table.rows[0]
    legends = [
        (NAVY, "All-team Phase"),
        (TEAL, "Lead Dev Track"),
        (PURPLE, "BA/PM Support Track"),
    ]
    for i, (color, label) in enumerate(legends):
        cell = row.cells[i]
        clear_cell(cell)
        set_cell_shading(cell, color)
        p = cell.paragraphs[0]
        run = p.add_run(label)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Word often duplicates merged cells — clear trailing duplicates
    for i in range(len(legends), len(row.cells)):
        clear_cell(row.cells[i])
        set_cell_shading(row.cells[i], WHITE)


def main():
    doc = Document(DOC_PATH)
    gantt = find_gantt_table(doc)
    fix_gantt_table(gantt)

    legend = find_legend_table(doc, gantt)
    if legend is not None:
        fix_legend_table(legend)

    doc.save(DOC_PATH)
    print(f"Gantt chart fixed: {DOC_PATH}")


if __name__ == "__main__":
    main()
